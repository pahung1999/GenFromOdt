from docx import Document
from docx.shared import Inches
import random
from src.odt_convert import *
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
import os
from tqdm import tqdm
from lorem_text import lorem
import subprocess
from src.get_label import labelme_gen

from docx2pdf import convert
import yaml
import json
import os
import fitz
from odf import opendocument
from src.odt_convert import replace_text, gen_info, gen_dump_text
from src.image_convert import bytes2pillow
from src.pdf_extraction import spiltpage_pdf
from src.get_label import labelme_gen
from src.table_extraction import extract_table, merge_table_element, column_from_text
from tqdm import tqdm
import subprocess
import camelot
from src.pdf_extraction import page_extraction

from src.docx_document import *

image_dir = "/home/phung/AnhHung/data/background/negative_labelme"
pdf_dir = "./output/pdf_files/"
docx_dir = "./output/docx_files/"
labelme_dir = "./output/labelme/"
label_list=['table', 'column', 'row', 'cell']

sample_num = 2


image_list = [os.path.join(image_dir, filename) for filename in os.listdir(image_dir) if "jpg" in filename]

def convert_docx_to_pdf(input_word_file, output_directory):
    command = [
        "libreoffice",
        "--headless",
        "--convert-to",
        "pdf",
        input_word_file,
        "--outdir",
        output_directory
    ]
    try:
        subprocess.run(command, check=True)
    except subprocess.CalledProcessError:
        print(f"Conversion failed: {input_word_file}")

print("Gen docx and convert to pdf...")
for file_num in tqdm(range(sample_num)):

    document = Document()

    # with open("./data/table_style_list.txt", "r") as f:
    #     table_styles = [x.replace('\n',"") for x in f.readlines()]
    # #Table config
    # # table_style = random.choice(table_styles)
    # table_style = table_styles[0]

    add_types = ["para", "pic", "table"]
    for i in range(1, 10):
        add_type = random.choice(add_types)
        if add_type == "para":
            document = add_para_to_doc(document)
        if add_type == "pic":
            document = add_pic_to_doc(document, image_list)
        if add_type == "table":
            document = add_table_to_doc(document, row_range = (5, 20), col_range = (2,6))
        document.add_paragraph("")

    # document.add_page_break()

    docx_path = os.path.join(docx_dir, f"{file_num:04d}.docx")
    document.save(docx_path)

    convert_docx_to_pdf(docx_path, pdf_dir)

print("Extract labelme from pdf...")
for pdf_file in tqdm(os.listdir(pdf_dir)):
    pdf_path = os.path.join(pdf_dir, pdf_file)
    pdf_doc = fitz.open(pdf_path)
    pdf_name, _ = os.path.splitext(pdf_file)

    for j, page in enumerate(pdf_doc):
        pixmap = page.get_pixmap()
        image = bytes2pillow(pixmap.tobytes())
        w, h = image.size

        texts, text_polygons = page_extraction(page, extract_type = 'word')

        shape_dict = {}
        tables = camelot.read_pdf(pdf_path, flavor = "lattice", pages=f'{j+1}')
        for table in tables:
            table_dict = extract_table(table, h)
            for table_key in table_dict:
                if table_key not in label_list:
                    continue
                if isinstance(table_dict[table_key][0] , list):
                    if table_key in shape_dict:
                        shape_dict[table_key] += table_dict[table_key]
                    else:
                        shape_dict[table_key] = table_dict[table_key]
                else:
                    if table_key in shape_dict:
                        shape_dict[table_key].append(table_dict[table_key])
                    else:
                        shape_dict[table_key] = [table_dict[table_key]]

        if 'table' in shape_dict:
            shape_dict['table'] = merge_table_element(shape_dict['table'], max_error = 10, box_type = 'table')
        
        if 'column' in shape_dict:
            shape_dict['column'] = merge_table_element(shape_dict['column'], max_error = 10, box_type = 'column')

        if 'row' in shape_dict:
            shape_dict['row'] = merge_table_element(shape_dict['row'], max_error = 10, box_type = 'row')

        if 'column' in shape_dict:
            shape_dict['column_text'] = column_from_text(shape_dict['column'], text_polygons)
        

        labelme_gen(shape_dict, 
                    image, 
                    f'{pdf_name}_page_{j:02d}', 
                    labelme_dir,
                    box_type = "rectangle")
        