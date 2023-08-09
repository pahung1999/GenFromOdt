from docx import Document
from docx.shared import Inches
import random
from src.odt_convert import *
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from lorem_text import lorem


def gen_dump_cols(row_len = 20, max_word = 6):
    return [lorem.words(random.randint(1, 3))] + [lorem.words(random.randint(1, max_word)) for i in range(row_len)]

def add_table_to_doc(document, row_range = (5, 20), col_range = (2,6), table_style = 'Table Grid'):
    
    rows = random.randint(*row_range)
    cols = random.randint(*col_range)
    fake = Faker()

    table = document.add_table(rows=rows, cols=cols)

    
    # with open("./data/table_style_list.txt", "r") as f:
    #     table_styles = [x.replace('\n',"") for x in f.readlines()]
    # #Table config
    # # table_style = random.choice(table_styles)
    # table_style = table_styles[file_num]
    if table_style is not None:
        table.style = document.styles[table_style] 
    table.autofit = True
    table.alignment = random.choice([WD_TABLE_ALIGNMENT.CENTER, WD_TABLE_ALIGNMENT.LEFT, WD_TABLE_ALIGNMENT.RIGHT])


    #Table content
    id_cols = [random.choice(["STT", "ID"])] + [f"{i+1:02d}" for i in range(rows)]
    phone_cols = [random.choice(["SĐT", "Phone number"])] + [random.choice([gen_code_string("NNNNNNNNNN"), gen_code_string("NNNNNNN-NNN"), gen_code_string("NNN NNNN NNN")]) for i in range(rows)]
    score_cols = [random.choice(["Điểm", "Score"])] + [str(random.randint(1,100)) for i in range(rows)]
    name_cols = [random.choice(["Họ và tên", "Last Name", "Full Name"])] + [gen_add_info("(name)", max_len=20, upper=False) for i in range(rows)]
    one_word_cols = [random.choice(["Tên", "Name"])] + [fake.word()[0].upper() + fake.word()[1:] for i in range(rows)]
    address_cols = [random.choice(["Địa chỉ", "Address"])] + [gen_add_info("(address)", max_len=80, upper=False) for i in range(rows)]
    country_cols = [random.choice(["Country", "Quốc tịch"])] + [gen_add_info("(country)", max_len=15, upper=False) for i in range(rows)]
    dump_words_cols = [generate_random_words(min_len=4, max_len=15)] + [generate_random_words(min_len=20, max_len=150) for i in range(rows)]
    type_cols = [random.choice(["Type", "Loại"])] + [random.choice(['A', 'B', 'C', 'D', 'E']) for i in range(rows)]

    full_table_info_dict = {
        "type1" : [id_cols, name_cols, one_word_cols, address_cols] + [gen_dump_cols(row_len = rows) for i in range(col_range[1]-4)],
        "type2" : [id_cols, name_cols, country_cols, score_cols, type_cols] + [gen_dump_cols(row_len = rows) for i in range(col_range[1]-5)],
        "type3" : [gen_dump_cols(row_len = rows, max_word = 20) for i in range(col_range[1])],
        "type4" : [id_cols, name_cols, one_word_cols, type_cols, phone_cols, address_cols] + [gen_dump_cols(row_len = rows) for i in range(col_range[1]-6)],
        "type5" : [id_cols, name_cols, address_cols] + [gen_dump_cols(row_len = rows) for i in range(col_range[1]-3)],
        "type6" : [name_cols, one_word_cols, type_cols] + [gen_dump_cols(row_len = rows) for i in range(col_range[1]-3)],
    }
    
    table_col_size_dict={
        "type1" : [1/14, 1/5, 1/10, 1/3.5] + [(1-sum([1/14, 1/5, 1/10, 1/3.5]))/(col_range[1]-4) for i in range(col_range[1]-4)],
        "type2" : [1/14, 1/4, 1/6, 1/12, 1/12] + [(1-sum([1/14, 1/4, 1/6, 1/12, 1/12]))/(col_range[1]-5) for i in range(col_range[1]-5)],
        "type3" : [1/col_range[1] for i in range(col_range[1])],
        "type4" : [1/14, 1/5, 1/10, 1/12, 1/6, 1/4] + [(1-sum([1/14, 1/5, 1/10, 1/12, 1/6, 1/4]))/(col_range[1]-6) for i in range(col_range[1]-6)],
        "type5" : [1/14, 1/5, 1/3.5] + [(1-sum([1/14, 1/5, 1/3.5]))/(col_range[1]-3) for i in range(col_range[1]-3)],
        "type6" : [1/5, 1/10, 1/12] + [(1-sum([1/5, 1/10, 1/12]))/(col_range[1]-3) for i in range(col_range[1]-3)],
    }

    random_key = random.choice(list(full_table_info_dict.keys()))
    full_table_info = full_table_info_dict[random_key]
    table_col_size = table_col_size_dict[random_key]
    if random_key == "type3":
        table_col_size= [1/(cols+0.5) for i in range(cols)]

    for row_id in range(rows):
        for col_id in range(cols):
            # try:
            table.rows[row_id].cells[col_id].text = full_table_info[col_id][row_id]
            table.rows[row_id].height_rule = WD_ROW_HEIGHT_RULE.AUTO
    
    for col_id in range(cols):
        table.columns[col_id].width = int(table_col_size[col_id]*5485130)
            
    # print("table_width = ", sum([table.columns[col_id].width for col_id in range(cols)]))
    return document

def add_para_to_doc(document):
    document.add_paragraph(lorem.words(random.randint(20, 150)), style = random.choice(['Normal', 'Body Text', 'Body Text 2', 'Body Text 3']))
    return document

def add_pic_to_doc(document, image_list):
    document.add_picture(random.choice(image_list), width=Inches(random.uniform(1, 3)))
    return document
